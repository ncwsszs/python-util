import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement

def set_font(run, font_name, size_pt, bold=False, color=(0,0,0)):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color)

def insert_paragraph_before(doc, text=None):
    # 在文档最前面插入一个空段落
    body = doc._body._element
    p = OxmlElement('w:p')  # 新段落元素
    body.insert(0, p)       # 插入到最前面
    new_para = Paragraph(p, doc)
    if text:
        new_para.add_run(text)
    return new_para

def set_paragraph_indent(paragraph, chars=2):
    # 首行缩进，1个字符约14.4pt
    paragraph.paragraph_format.first_line_indent = Pt(14.4 * chars)

def set_table_border(table):
    for cell in table._cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for border_name in ['top','left','bottom','right','insideH','insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            borders.append(border)
        tcPr.append(borders)

def create_styles(doc):
    styles = doc.styles

    # 标题1（大标题）
    if '自定义标题1' not in styles:
        style = styles.add_style('自定义标题1', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = '仿宋'
        font.size = Pt(26)  # 一号
        font.bold = True
        font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 标题2（模块标题）
    if '自定义标题2' not in styles:
        style = styles.add_style('自定义标题2', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = '仿宋'
        font.size = Pt(16)  # 三号
        font.bold = True
        font.color.rgb = RGBColor(0, 0, 0)
        pf = style.paragraph_format
        pf.left_indent = Pt(0)
        pf.first_line_indent = None
        pf.space_before = Pt(13)
        pf.space_after = Pt(13)
        pf.line_spacing = Pt(24)

    # 正文（my正文）
    if 'my正文' not in styles:
        style = styles.add_style('my正文', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = '仿宋'
        font.size = Pt(14)  # 四号
        font.bold = False
        font.color.rgb = RGBColor(0, 0, 0)
        pf = style.paragraph_format
        pf.first_line_indent = Pt(28)  # 两字符缩进
        pf.line_spacing = 1.5


def set_table_border(table):
    tbl = table._element
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set('w:val', 'single')
        border.set('w:sz', '4')
        border.set('w:space', '0')
        border.set('w:color', '000000')
        tblBorders.append(border)
    tbl.tblPr.append(tblBorders)


def format_doc(doc, filename):
    create_styles(doc)

    # 添加大标题
    title_para = doc.paragraphs[0].insert_paragraph_before(filename)
    title_para.style = '自定义标题1'

    module_titles = [
        '简要描述', '请求URL', '请求方式', 'Header',
        '请求参数', '返回示例', '返回参数说明', '备注'
    ]

    # 处理段落
    module_counter = 1
    for para in doc.paragraphs[1:]:
        text = para.text.strip()
        if not text:
            continue
        if any(text == t or text.startswith(t) for t in module_titles):
            para.style = '自定义标题2'
            para.text = f"{module_counter}、{text}"
            module_counter += 1
        else:
            para.style = 'my正文'

    # 处理表格
    for table in doc.tables:
        set_table_border(table)
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.style = 'my正文'
                    para.paragraph_format.first_line_indent = None
                    for run in para.runs:
                        run.font.size = Pt(10.5)  # 五号
                        run.font.name = '仿宋'
                        run.font.color.rgb = RGBColor(0, 0, 0)



def process_file(input_path, output_path):
    doc = Document(input_path)
    filename = os.path.splitext(os.path.basename(input_path))[0]
    format_doc(doc, filename)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"已处理: {output_path}")

def batch_process_folder(input_folder, output_folder):
    for root, _, files in os.walk(input_folder):
        for file in files:
            if file.lower().endswith('.docx'):
                input_path = os.path.join(root, file)
                rel_path = os.path.relpath(input_path, input_folder)
                output_path = os.path.join(output_folder, rel_path)
                process_file(input_path, output_path)

if __name__ == "__main__":
    input_folder = r"G:\OneDrive\文档\工作\瑞景中台\对接文档\word\智慧生态园林对接"
    output_folder = r"G:\OneDrive\文档\工作\瑞景中台\对接文档\word\智慧生态园林对接2"
    batch_process_folder(input_folder, output_folder)
