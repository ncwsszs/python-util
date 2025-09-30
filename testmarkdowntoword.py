# -*- coding: utf-8 -*-
import os
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown import markdown
from bs4 import BeautifulSoup

def clean_empty_headings(md_text: str) -> str:
    md_text = md_text.replace('\u00A0', ' ')
    md_text = md_text.replace('\u200B', '')
    md_text = md_text.replace('\ufeff', '')

    md_text = re.sub(
        r'^[ \t\u00A0\u200B\uFEFF]*#{1,6}[ \t\u00A0\u200B\uFEFF]*\r?\n',
        '',
        md_text,
        flags=re.MULTILINE
    )
    md_text = re.sub(
        r'^[ \t\u00A0\u200B\uFEFF]*#{1,6}[ \t\u00A0\u200B\uFEFF]*$',
        '',
        md_text,
        flags=re.MULTILINE
    )

    md_text = re.sub(r'\n{3,}', '\n\n', md_text)
    md_text = md_text.strip() + '\n'
    return md_text

def find_style_by_substr(doc: Document, substr: str):
    if not substr:
        return None
    for s in doc.styles:
        if s.name and substr in s.name:
            return s.name
    return None

def set_cell_border(cell, color="000000", size="4"):
    """为单元格设置边框（size单位: 八分之一磅）"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for border_name in ("top", "left", "bottom", "right"):
        element = tcPr.xpath(f"./w:{border_name}")[0] if tcPr.xpath(f"./w:{border_name}") else None
        if not element:
            from docx.oxml import OxmlElement
            element = OxmlElement(f"w:{border_name}")
            tcPr.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)
        element.set(qn("w:space"), "0")

def html_to_docx_into(doc: Document, html: str, file_title: str,
                      body_style_hint='my正文', table_style_hint='表格字体样式',
                      heading_style_candidates=None):
    soup = BeautifulSoup(html, 'html.parser')

    # 插入标题
    p = doc.add_paragraph()
    run = p.add_run(file_title)
    run.font.name = '仿宋'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    run.font.size = Pt(26)
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中

    body_style = find_style_by_substr(doc, body_style_hint) or 'Normal'
    table_style = find_style_by_substr(doc, table_style_hint)

    if heading_style_candidates is None:
        heading_style_candidates = ['Heading 2', '标题 2', '标题2']
    heading_style = None
    for cand in heading_style_candidates:
        heading_style = find_style_by_substr(doc, cand)
        if heading_style:
            break

    top_elems = [e for e in soup.contents if not (isinstance(e, str) and e.strip() == '')]
    for elem in top_elems:
        name = getattr(elem, 'name', None)
        if name and name.lower() in ('h1','h2','h3','h4','h5','h6'):
            text = elem.get_text().strip()
            if not text:
                continue
            doc.add_paragraph(text, style=heading_style or 'Heading 2')
        elif name == 'p' or (not name and elem.strip() != ''):
            text = elem.get_text().strip()
            para = doc.add_paragraph(text, style=body_style)
            para.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进两字符
        elif name in ('ul', 'ol'):
            # 列表改为普通段落+首行缩进，不加点
            for li in elem.find_all('li', recursive=False):
                text = li.get_text().strip()
                para = doc.add_paragraph(text, style=body_style)
                para.paragraph_format.first_line_indent = Cm(0.74)
        elif name == 'table':
            rows = elem.find_all('tr')
            if not rows:
                continue
            ncols = max(len(r.find_all(['th','td'])) for r in rows)
            table = doc.add_table(rows=len(rows), cols=ncols)
            for i, row in enumerate(rows):
                cells = row.find_all(['th','td'])
                for j in range(ncols):
                    target_cell = table.cell(i, j)
                    if j < len(cells):
                        cell_text = cells[j].get_text().strip()
                    else:
                        cell_text = ''
                    target_para = target_cell.paragraphs[0]
                    target_para.text = cell_text
                    if table_style:
                        try:
                            target_para.style = table_style
                        except KeyError:
                            pass
                    set_cell_border(target_cell)  # 设置黑色全边框
        else:
            text = elem.get_text().strip()
            if text:
                para = doc.add_paragraph(text, style=body_style)
                para.paragraph_format.first_line_indent = Cm(0.74)

def convert_markdown_to_docx(md_path, docx_path, template_file):
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    md_content = clean_empty_headings(md_content)
    html = markdown(md_content, extensions=['extra', 'tables', 'fenced_code'])

    if not template_file or not os.path.exists(template_file):
        raise FileNotFoundError(f"模板文件不存在：{template_file}")
    doc = Document(template_file)

    title_text = os.path.splitext(os.path.basename(md_path))[0]
    html_to_docx_into(doc, html, title_text,
                      body_style_hint='1.my正文',
                      table_style_hint='2.表格字体样式')

    os.makedirs(os.path.dirname(docx_path), exist_ok=True)
    doc.save(docx_path)

def batch_convert_folder(src_folder, dst_folder, template_file):
    for root, _, files in os.walk(src_folder):
        for file in files:
            if file.lower().endswith('.md'):
                md_path = os.path.join(root, file)
                rel_path = os.path.relpath(md_path, src_folder)
                rel_path_no_ext = os.path.splitext(rel_path)[0]
                docx_path = os.path.join(dst_folder, rel_path_no_ext + '.docx')
                print(f"转换: {md_path} -> {docx_path}")
                convert_markdown_to_docx(md_path, docx_path, template_file)

if __name__ == '__main__':
    src = r"markdown"
    dst = r"word2"
    template = r"一般文档模板.docx"
    batch_convert_folder(src, dst, template)
